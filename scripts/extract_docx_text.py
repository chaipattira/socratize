#!/usr/bin/env python3
"""Extract text from a .docx file using python-docx, output to stdout."""
import sys
from docx import Document


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: extract_docx_text.py <docx_file>", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    try:
        doc = Document(path)
    except Exception as e:
        print(f"Error reading docx: {e}", file=sys.stderr)
        sys.exit(1)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append("\t".join(cells))

    if not paragraphs:
        print("(no text content found in document)", file=sys.stderr)
        sys.exit(1)

    print("\n".join(paragraphs))


if __name__ == "__main__":
    main()
